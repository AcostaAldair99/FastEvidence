##This is de code for every evidence (docx file)
from docx import Document
from docx.shared import Inches 
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import ImageGrab
import datetime
import os

class Evidence:

    def __init__(self,path,filename,metadata,description):
      self.path = path
      self.filename = filename
      self.executionData= metadata 
      self.doc = ""
      self.executionDate = datetime.datetime.today().strftime("%d-%m-%Y")
      self.clockTime = datetime.datetime.now().strftime("%H:%M:%S")
      self.description = description
      self.dir = self.path +"/"+ self.filename#str(random.randint(1,200))
      self.picture_dir = self.dir + "/pictures" 
      self.pictures_num = 0    

    def toString(self):
      print("Path: "+self.path)
      print("Filename: "+self.filename)
      print("Dir: "+self.dir)


    def createDirectories(self):
      if os.path.exists(self.dir):
        return True
      else:
        os.mkdir(self.dir)      
        os.mkdir(self.picture_dir)
        return False
    

    def closeDocument(self):
      dir = str(self.dir+"/"+self.filename+".docx")
      self.doc.save(dir)

    def createDocument(self):
      self.doc = Document()
      self.setHeader()
      self.settingTable()
      self.setTestDescription()


    def setHeader(self):
      section = self.doc.sections[0]
      header = section.header
      paragraph = header.paragraphs[0]
      run = paragraph.add_run(self.clockTime+"\t\tQA EXECUTION EVIDENCE [ "+self.executionDate+" ]")
      run.bold = True 
      run.italic = True
      run.font.size = Pt(10)
      paragraph.style = self.doc.styles["Header"]
    
    
    def setCellStyle(self,cell,defect):
      paragraph = cell.paragraphs[0]
      run = paragraph.runs[0]
      run.bold = True
      if defect == True:
        shadding_elm = parse_xml(r'<w:shd {} w:fill="FF4D4D"/>'.format(nsdecls('w')))
      else:
        shadding_elm = parse_xml(r'<w:shd {} w:fill="80AAFF"/>'.format(nsdecls('w')))
      cell._tc.get_or_add_tcPr().append(shadding_elm)

    def setCellStatusStyle(self,cell,value):
      paragraph = cell.paragraphs[0]
      run = paragraph.runs[0]
      run.bold = True
      if value == "PASSED":
        shadding_elm = parse_xml(r'<w:shd {} w:fill="FF4D4D"/>'.format(nsdecls('w')))
      elif value == "FAILED":
        shadding_elm = parse_xml(r'<w:shd {} w:fill="BFFF80"/>'.format(nsdecls('w')))
      else:
        shadding_elm = parse_xml(r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w')))

      cell._tc.get_or_add_tcPr().append(shadding_elm)

    def settingTable(self):
      table_headers = ("dInterface","Request","Test-Case","Status","Tester","Defect")    
      table = self.doc.add_table(rows = 0,cols = 2)
      table.columns[0].width = Inches(1.0)
      table.columns[1].width = Inches(6.0)
      for header in table_headers:
        cell = table.add_row().cells
        color = False
        cell[0].text = header 
        if header == "Defect":
          color = True
        self.setCellStyle(cell[0],color)

      for i, row in enumerate(table.rows):
        cell = row.cells[1]
        if i < len(self.executionData):
          #To update -> the code to highlight the background color in function of if the test case is passed or not
          # self.setCellStatusStyle(cell[1],self.executionData[i])
          cell.text = self.executionData[i] 

      self.doc.add_paragraph()
      
    def setTestDescription(self):
      paragraph = self.doc.add_paragraph()
      paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
      run = paragraph.add_run("Description: "+self.description)
      run.bold = True
      run.font.size = Pt(14)
      run.italic = True
      self.doc.add_paragraph()

    def addPicture(self,description):
      screenshot = self.getScreenshot()
      self.addStepDescription(description)
      paragraph = self.doc.add_paragraph()
      run = paragraph.add_run()
      run.add_picture(screenshot, width = Inches(7.0))
      paragraph.aligment = 1

    def addStepDescription(self,description):
      paragraph = self.doc.add_paragraph()
      run = paragraph.add_run(str(self.pictures_num)+" : "+description)
      run.bold = True
      run.font.size = Pt(12)

    def getScreenshot(self):
      screenshot = ImageGrab.grab()
      self.pictures_num+=1
      filename = "Step-"+str(self.pictures_num)+".png"
      screenshot.save(self.picture_dir+"/"+filename)
      return self.picture_dir+"/"+filename

